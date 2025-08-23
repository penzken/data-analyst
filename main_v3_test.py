import requests
import json
import os
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from dotenv import load_dotenv
from typing import TypedDict, List, Dict, Any, Optional
from langgraph.graph import StateGraph, END
from langchain_core.messages import SystemMessage, HumanMessage
from langchain_core.tools import tool
from langchain_google_genai import ChatGoogleGenerativeAI
from datetime import datetime
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches

# Import các hàm từ các file khác
from preprocessing import preprocessing_data, safe_convert_to_number
from utils.knowledge.knowledge_base import retrieve_knowledge # Mới

# Tải biến môi trường (bao gồm cả cấu hình LangSmith)
load_dotenv(os.path.join(os.path.dirname(__file__), '.env'))

# --- Cấu hình ---
N8N_URL = 'https://aizen.auto.123host.asia/webhook/9057a9e7-1f29-4474-a501-730a2f2bfb68'
CHARTS_DIR = 'charts'
REPORTS_DIR = 'reports'
MAX_REFLECTIONS = 3 # Giới hạn số lần tự cải thiện để tránh lặp vô hạn

os.makedirs(CHARTS_DIR, exist_ok=True)
os.makedirs(REPORTS_DIR, exist_ok=True)

def extract_json_from_string(text: str) -> Optional[Dict]:
    """
    Trích xuất một đối tượng JSON từ một chuỗi, ngay cả khi nó được bao quanh bởi văn bản khác.
    """
    try:
        # Tìm vị trí của dấu '{' đầu tiên và dấu '}' cuối cùng
        start_index = text.find('{')
        end_index = text.rfind('}')
        
        if start_index != -1 and end_index != -1 and end_index > start_index:
            json_str = text[start_index:end_index+1]
            return json.loads(json_str)
        else:
            return None
    except (json.JSONDecodeError, IndexError):
        return None
    
# --- State của LangGraph (Mở rộng) ---
class Reflection:
    def __init__(self, analysis: str, critique: str):
        self.analysis = analysis
        self.critique = critique

class AnalysisState(TypedDict, total=False):
    question: str
    payload: Any
    rows: List[Dict[str, Any]]
    context: str
    analysis: str
    calculations: Dict[str, Any]
    chart_paths: Dict[str, str]
    report_path: str
    knowledge: List[str] # Mới: Lưu kiến thức được truy xuất
    reflection_history: List[Reflection] # Mới: Lịch sử các lần tự cải thiện
    current_score: int

# --- Các hàm tính toán (Giữ nguyên từ file của bạn) ---
@tool
def calculate_total_revenue(rows: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Calculate total revenue from order data"""
    total_revenue = 0
    unique_orders = set()
    
    for row in rows:
        order_id = row.get('orderId')
        calc_total = safe_convert_to_number(row.get('calcTotalMoney'))
        
        # Only count each order once for total revenue
        if order_id and order_id not in unique_orders:
            unique_orders.add(order_id)
            total_revenue += calc_total
    
    return {
        "total_revenue": total_revenue,
        "total_orders": len(unique_orders),
        "average_order_value": total_revenue / len(unique_orders) if unique_orders else 0
    }

@tool
def calculate_product_stats(rows: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Calculate product statistics including top products by quantity and revenue"""
    product_stats = {}
    
    for row in rows:
        product_name = row.get('productName', 'Unknown')
        quantity = safe_convert_to_number(row.get('quantity'))
        price = safe_convert_to_number(row.get('price'))
        
        if product_name not in product_stats:
            product_stats[product_name] = {
                'total_quantity': 0,
                'total_revenue': 0,
                'order_count': 0
            }
        
        product_stats[product_name]['total_quantity'] += quantity
        product_stats[product_name]['total_revenue'] += quantity * price
        product_stats[product_name]['order_count'] += 1
    
    # Sort by quantity and revenue
    top_by_quantity = sorted(
        product_stats.items(), 
        key=lambda x: x[1]['total_quantity'], 
        reverse=True
    )[:5]
    
    top_by_revenue = sorted(
        product_stats.items(), 
        key=lambda x: x[1]['total_revenue'], 
        reverse=True
    )[:5]
    
    return {
        "top_products_by_quantity": [
            {
                "product_name": name,
                "total_quantity": stats['total_quantity'],
                "total_revenue": stats['total_revenue'],
                "order_count": stats['order_count']
            }
            for name, stats in top_by_quantity
        ],
        "top_products_by_revenue": [
            {
                "product_name": name,
                "total_quantity": stats['total_quantity'],
                "total_revenue": stats['total_revenue'],
                "order_count": stats['order_count']
            }
            for name, stats in top_by_revenue
        ]
    }

@tool
def calculate_daily_stats(rows: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Calculate daily statistics from order data"""
    daily_stats = {}
    
    for row in rows:
        date = row.get('date', '')
        if not date:
            continue
        
        quantity = safe_convert_to_number(row.get('quantity'))
        price = safe_convert_to_number(row.get('price'))
        revenue = quantity * price
        
        if date not in daily_stats:
            daily_stats[date] = {
                'total_quantity': 0,
                'total_revenue': 0,
                'unique_orders': set(),
                'unique_products': set()
            }
        
        daily_stats[date]['total_quantity'] += quantity
        daily_stats[date]['total_revenue'] += revenue
        daily_stats[date]['unique_orders'].add(row.get('orderId'))
        daily_stats[date]['unique_products'].add(row.get('productName'))
    
    # Convert sets to counts
    for date_stats in daily_stats.values():
        date_stats['order_count'] = len(date_stats['unique_orders'])
        date_stats['product_count'] = len(date_stats['unique_products'])
        del date_stats['unique_orders']
        del date_stats['unique_products']
    
    return daily_stats

@tool
def calculate_hourly_stats(rows: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Calculate hourly statistics from order data"""
    hourly_stats = {}
    time_period_stats = {}
    
    for row in rows:
        hour = row.get('hour', '')
        time_period = row.get('time_period', '')
        
        if not hour:
            continue
        
        quantity = safe_convert_to_number(row.get('quantity'))
        price = safe_convert_to_number(row.get('price'))
        revenue = quantity * price
        
        # Hourly stats
        if hour not in hourly_stats:
            hourly_stats[hour] = {
                'total_quantity': 0,
                'total_revenue': 0,
                'unique_orders': set(),
                'unique_products': set()
            }
        
        hourly_stats[hour]['total_quantity'] += quantity
        hourly_stats[hour]['total_revenue'] += revenue
        hourly_stats[hour]['unique_orders'].add(row.get('orderId'))
        hourly_stats[hour]['unique_products'].add(row.get('productName'))
        
        # Time period stats
        if time_period and time_period not in time_period_stats:
            time_period_stats[time_period] = {
                'total_quantity': 0,
                'total_revenue': 0,
                'unique_orders': set(),
                'unique_products': set()
            }
        
        if time_period:
            time_period_stats[time_period]['total_quantity'] += quantity
            time_period_stats[time_period]['total_revenue'] += revenue
            time_period_stats[time_period]['unique_orders'].add(row.get('orderId'))
            time_period_stats[time_period]['unique_products'].add(row.get('productName'))
    
    # Convert sets to counts and sort
    for hour_stats in hourly_stats.values():
        hour_stats['order_count'] = len(hour_stats['unique_orders'])
        hour_stats['product_count'] = len(hour_stats['unique_products'])
        del hour_stats['unique_orders']
        del hour_stats['unique_products']
    
    for period_stats in time_period_stats.values():
        period_stats['order_count'] = len(period_stats['unique_orders'])
        period_stats['product_count'] = len(period_stats['unique_products'])
        del period_stats['unique_orders']
        del period_stats['unique_products']
    
    # Find busiest hours and periods
    busiest_hours = sorted(
        hourly_stats.items(),
        key=lambda x: x[1]['order_count'],
        reverse=True
    )[:5]
    
    busiest_periods = sorted(
        time_period_stats.items(),
        key=lambda x: x[1]['order_count'],
        reverse=True
    )
    
    return {
        "hourly_breakdown": hourly_stats,
        "time_period_breakdown": time_period_stats,
        "busiest_hours": [
            {
                "hour": hour,
                "order_count": stats['order_count'],
                "total_revenue": stats['total_revenue'],
                "total_quantity": stats['total_quantity']
            }
            for hour, stats in busiest_hours
        ],
        "busiest_periods": [
            {
                "period": period,
                "order_count": stats['order_count'],
                "total_revenue": stats['total_revenue'],
                "total_quantity": stats['total_quantity']
            }
            for period, stats in busiest_periods
        ]
    }

def calculate_comprehensive_stats(rows: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Calculate all statistics using the calculator tools"""
    revenue_stats = calculate_total_revenue.invoke({"rows": rows})
    product_stats = calculate_product_stats.invoke({"rows": rows})
    daily_stats = calculate_daily_stats.invoke({"rows": rows})
    hourly_stats = calculate_hourly_stats.invoke({"rows": rows})
    
    return {
        "revenue_summary": revenue_stats,
        "product_analysis": product_stats,
        "daily_breakdown": daily_stats,
        "time_analysis": hourly_stats,
        "data_quality": {
            "total_rows": len(rows),
            "rows_with_valid_price": sum(1 for row in rows if safe_convert_to_number(row.get('price')) > 0),
            "rows_with_valid_quantity": sum(1 for row in rows if safe_convert_to_number(row.get('quantity')) > 0),
            "rows_with_valid_time": sum(1 for row in rows if row.get('hour', '') != ''),
        }
    }

# --- Các hàm Node của LangGraph ---
# Hàm tiện ích để fetch data (giữ nguyên)
def fetch_data(from_date: str, to_date: str) -> Dict[str, Any]:
    """Fetch data from n8n webhook"""
    data = {
        "fromDate": from_date,
        "toDate": to_date
    }
    headers = {'Content-Type': 'application/json'}
    response = requests.post(N8N_URL, data=json.dumps(data), headers=headers)
    response.raise_for_status()
    return response.json()

def preprocess_node(state: AnalysisState) -> Dict[str, Any]:
    """Node tiền xử lý dùng payload đã được truyền vào (không gọi API)."""
    print("🔄 Bắt đầu tiền xử lý và tính toán từ payload đã cung cấp...")
    payload = state.get("payload")
    if payload is None:
        raise ValueError("Thiếu payload đầu vào. Hãy truyền payload từ N8N vào app.invoke.")
    
    rows = preprocessing_data(payload)
    calculations = calculate_comprehensive_stats(rows)
    
    context_data = {
        "raw_data_sample": rows[:5],
        "total_rows": len(rows),
        "calculations": calculations
    }
    context_json = json.dumps(context_data, ensure_ascii=False, default=str)
    
    print("✅ Tiền xử lý hoàn tất.")
    return {"payload": payload, "rows": rows, "context": context_json, "calculations": calculations}

def retrieve_knowledge_node(state: AnalysisState) -> Dict[str, Any]:
    """Node để truy xuất kiến thức liên quan."""
    print("🧠 Đang truy xuất kiến thức từ cơ sở tri thức...")
    question = state.get("question", "")
    knowledge = retrieve_knowledge(question)
    return {"knowledge": knowledge, "reflection_history": []}

def llm_analyst_node(state: AnalysisState) -> Dict[str, Any]:
    """Node LLM để tạo bản phân tích (bản nháp)."""
    print("✍️ LLM Analyst đang tạo bản nháp báo cáo...")
    question = state["question"]
    context = state["context"]
    knowledge = state["knowledge"]
    history = state.get("reflection_history", [])
    calculations = state.get("calculations", {})
    
    google_api_key = os.getenv("GOOGLE_API_KEY")
    llm = ChatGoogleGenerativeAI(model="gemini-2.5-flash", temperature=0.3, api_key=google_api_key)
    
    knowledge_prompt = "\n\n**Gợi ý từ chuyên gia (Mentor):**\n" + "\n".join(f"- {k}" for k in knowledge) if knowledge else ""
    
    history_prompt = ""
    if history:
        last_critique = history[-1]['critique']
        history_prompt = f"\n\n**Phản hồi từ lần trước (cần cải thiện):**\n{last_critique}\nHãy viết lại báo cáo dựa trên phản hồi này."

    calc_summary = json.dumps(calculations, ensure_ascii=False, indent=2, default=str)

    messages = [
        SystemMessage(content=(
            "Bạn là chuyên gia phân tích dữ liệu ngành F&B. "
            "Nhiệm vụ của bạn là tạo ra một báo cáo phân tích chi tiết, chuyên nghiệp từ các số liệu đã được tính toán sẵn. "
            "Sử dụng định dạng markdown và tập trung đưa ra các nhận định (insights) và đề xuất (recommendations) actionable. "
            "Trả lời bằng tiếng Việt."
        )),
        HumanMessage(content=(
            f"Dữ liệu và kết quả tính toán:\n{context}\n\n"
            f"Các số liệu đã được tính toán chính xác:\n{calc_summary}\n\n"
            f"Câu hỏi: {question}\n\n"
            "Yêu cầu: Dựa trên các số liệu đã tính toán, hãy tạo báo cáo tóm tắt bao gồm:\n"
            "1. Tổng quan (số đơn hàng, tổng doanh thu, giá trị đơn hàng trung bình)\n"
            "2. Top 5 sản phẩm theo số lượng bán\n"
            "3. Top 5 sản phẩm theo doanh thu\n"
            "4. Phân tích theo ngày (nếu có)\n"
            "5. Khung giờ nào đông khách nhất\n"
            "6. Đánh giá chất lượng dữ liệu\n"
            "7. Insight và khuyến nghị\n\n"
            "Sử dụng định dạng markdown để trình bày báo cáo một cách rõ ràng và chuyên nghiệp."
        )),
    ]
    
    response = llm.invoke(messages)
    print("✅ LLM Analyst đã hoàn thành bản nháp.")
    return {"analysis": response.content}

def llm_critic_node(state: AnalysisState) -> Dict[str, Any]:
    """Node LLM Critic để đánh giá bản phân tích (phiên bản nâng cao)."""
    print("🧐 LLM Critic đang đánh giá báo cáo...")
    analysis = state["analysis"]
    knowledge = state["knowledge"]
    
    google_api_key = os.getenv("GOOGLE_API_KEY")
    # Bật chế độ JSON để tăng độ tin cậy
    llm = ChatGoogleGenerativeAI(
        model="gemini-2.5-flash", # Sửa lại tên model chính xác
        temperature=0.2, 
        api_key=google_api_key,
        model_kwargs={"response_format": {"type": "json_object"}}
    )
    
    knowledge_prompt = "\n\n**Các tiêu chí cần tham khảo từ Mentor:**\n" + "\n".join(f"- {k}" for k in knowledge) if knowledge else ""

    messages = [
        SystemMessage(content=(
            "Bạn là một người quản lý/chuyên gia phân tích dữ liệu F&B cực kỳ khó tính. "
            "Nhiệm vụ của bạn là đánh giá một báo cáo và trả về một đối tượng JSON duy nhất. "
            "Không thêm bất kỳ văn bản nào khác ngoài đối tượng JSON."
        )),
        HumanMessage(content=(
            f"Đây là báo cáo cần đánh giá:\n\n---\n{analysis}\n---\n"
            f"{knowledge_prompt}"
            "\n\nYêu cầu:\n"
            "1. Đánh giá báo cáo trên theo thang điểm 10 (về mức độ sâu sắc và tính hữu dụng).\n"
            "2. Nếu điểm dưới 8, hãy đưa ra những phản hồi cụ thể, rõ ràng để nhân viên có thể cải thiện.\n"
            "3. Trả lời chỉ bằng một chuỗi JSON theo định dạng sau: "
            '{"score": <điểm_số>, "critique": "<nội_dung_phản_hồi>"}'
        )),
    ]
    
    response = llm.invoke(messages)
    
    # Sử dụng hàm trích xuất JSON thông minh
    result = extract_json_from_string(response.content)
    
    if result and 'score' in result and 'critique' in result:
        score = result.get("score", 0)
        critique = result.get("critique", "Không có phản hồi.")
        print(f"✅ LLM Critic đã đánh giá: Điểm {score}/10.")
    else:
        print(" Lỗi: Critic trả về định dạng không hợp lệ hoặc thiếu trường. Gán điểm mặc định.")
        score = 0 # Gán điểm thấp để yêu cầu làm lại
        critique = "Phản hồi từ Critic không đúng định dạng. Yêu cầu viết lại báo cáo rõ ràng hơn."

    history = state.get("reflection_history", [])
    history.append(Reflection(analysis=analysis, critique=critique))
    
    # Cập nhật cả lịch sử và điểm số hiện tại
    return {"reflection_history": history, "current_score": score}

# --- Logic điều kiện cho Graph ---
def should_continue(state: AnalysisState) -> str:
    """Quyết định xem nên kết thúc hay cần cải thiện báo cáo."""
    history = state.get("reflection_history", [])
    score = state.get("current_score", 0)

    # Nếu số lần cải thiện vượt quá giới hạn -> kết thúc
    if len(history) >= MAX_REFLECTIONS: # Sửa thành >= để logic đúng hơn
        print("⚠️ Đạt giới hạn số lần tự cải thiện. Chấp nhận phiên bản cuối cùng.")
        return "end"

    if score >= 8:
        print("👍 Báo cáo được chấp nhận. Chuyển sang các bước tiếp theo.")
        return "end"
    else:
        print("👎 Báo cáo cần cải thiện. Gửi lại cho Analyst.")
        return "reflect"
    
# --- Các node còn lại (Tạo biểu đồ, báo cáo, gửi email) ---
def create_charts_node(state: AnalysisState) -> Dict[str, Any]:
    """Node để tạo biểu đồ từ dữ liệu."""
    print("📊 Đang tạo biểu đồ...")
    rows = state["rows"]
    calculations = state["calculations"]
    
    chart_paths = {}
    
    try:
        # Tạo biểu đồ doanh thu theo ngày
        if "daily_breakdown" in calculations:
            daily_data = calculations["daily_breakdown"]
            dates = list(daily_data.keys())
            revenues = [daily_data[date]["total_revenue"] for date in dates]
            
            plt.figure(figsize=(12, 6))
            plt.plot(dates, revenues, marker='o', linewidth=2, markersize=8)
            plt.title("Doanh Thu Theo Ngày", fontsize=16, fontweight='bold')
            plt.xlabel("Ngày", fontsize=12)
            plt.ylabel("Doanh Thu (VNĐ)", fontsize=12)
            plt.xticks(rotation=45)
            plt.grid(True, alpha=0.3)
            plt.tight_layout()
            
            chart_path = os.path.join(CHARTS_DIR, "daily_revenue.png")
            plt.savefig(chart_path, dpi=300, bbox_inches='tight')
            plt.close()
            chart_paths["daily_revenue"] = chart_path
        
        # Tạo biểu đồ top sản phẩm
        if "product_analysis" in calculations:
            product_data = calculations["product_analysis"]["top_products_by_revenue"][:5]
            products = [item["product_name"] for item in product_data]
            revenues = [item["total_revenue"] for item in product_data]
            
            plt.figure(figsize=(10, 6))
            bars = plt.bar(products, revenues, color='skyblue', alpha=0.8)
            plt.title("Top 5 Sản Phẩm Theo Doanh Thu", fontsize=16, fontweight='bold')
            plt.xlabel("Sản Phẩm", fontsize=12)
            plt.ylabel("Doanh Thu (VNĐ)", fontsize=12)
            plt.xticks(rotation=45, ha='right')
            plt.grid(True, alpha=0.3, axis='y')
            
            # Thêm giá trị trên mỗi cột
            for bar, revenue in zip(bars, revenues):
                plt.text(bar.get_x() + bar.get_width()/2, bar.get_height() + max(revenues)*0.01,
                        f'{revenue:,.0f}', ha='center', va='bottom', fontweight='bold')
            
            plt.tight_layout()
            
            chart_path = os.path.join(CHARTS_DIR, "top_products.png")
            plt.savefig(chart_path, dpi=300, bbox_inches='tight')
            plt.close()
            chart_paths["top_products"] = chart_path
        
        print("✅ Tạo biểu đồ hoàn tất.")
        return {"chart_paths": chart_paths}
        
    except Exception as e:
        print(f"❌ Lỗi khi tạo biểu đồ: {e}")
        return {"chart_paths": {}}

def create_report_node(state: AnalysisState) -> Dict[str, Any]:
    """Node để tạo báo cáo Word."""
    print("📝 Đang tạo báo cáo Word...")
    analysis = state["analysis"]
    calculations = state["calculations"]
    chart_paths = state.get("chart_paths", {})
    
    try:
        doc = Document()
        
        # Tiêu đề
        title = doc.add_heading("Báo Cáo Phân Tích Dữ Liệu F&B", 0)
        title.alignment = 1  # Căn giữa
        
        # Thông tin thời gian
        doc.add_paragraph(f"Ngày tạo: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
        doc.add_paragraph("")
        
        # Tóm tắt
        doc.add_heading("Tóm Tắt", level=1)
        summary = doc.add_paragraph()
        summary.add_run("Báo cáo này trình bày kết quả phân tích dữ liệu kinh doanh ngành F&B, bao gồm:")
        summary.add_run("\n• Phân tích doanh thu và đơn hàng")
        summary.add_run("\n• Top sản phẩm bán chạy")
        summary.add_run("\n• Phân tích theo thời gian")
        summary.add_run("\n• Các khuyến nghị kinh doanh")
        
        # Thống kê tổng quan
        doc.add_heading("Thống Kê Tổng Quan", level=1)
        if "revenue_summary" in calculations:
            revenue_data = calculations["revenue_summary"]
            doc.add_paragraph(f"• Tổng doanh thu: {revenue_data.get('total_revenue', 0):,.0f} VNĐ")
            doc.add_paragraph(f"• Tổng số đơn hàng: {revenue_data.get('total_orders', 0):,}")
            doc.add_paragraph(f"• Giá trị đơn hàng trung bình: {revenue_data.get('average_order_value', 0):,.0f} VNĐ")
        
        # Phân tích sản phẩm
        doc.add_heading("Phân Tích Sản Phẩm", level=1)
        if "product_analysis" in calculations:
            product_data = calculations["product_analysis"]
            
            doc.add_heading("Top 5 Sản Phẩm Theo Số Lượng", level=2)
            for i, product in enumerate(product_data.get("top_products_by_quantity", [])[:5], 1):
                doc.add_paragraph(f"{i}. {product['product_name']}: {product['total_quantity']:,} đơn vị")
            
            doc.add_heading("Top 5 Sản Phẩm Theo Doanh Thu", level=2)
            for i, product in enumerate(product_data.get("top_products_by_revenue", [])[:5], 1):
                doc.add_paragraph(f"{i}. {product['product_name']}: {product['total_revenue']:,.0f} VNĐ")
        
        # Phân tích theo thời gian
        doc.add_heading("Phân Tích Theo Thời Gian", level=1)
        if "time_analysis" in calculations:
            time_data = calculations["time_analysis"]
            
            doc.add_heading("Giờ Cao Điểm", level=2)
            for hour_data in time_data.get("busiest_hours", [])[:5]:
                doc.add_paragraph(f"• Giờ {hour_data['hour']}: {hour_data['order_count']} đơn hàng")
        
        # Thêm biểu đồ
        if chart_paths:
            doc.add_heading("Biểu Đồ", level=1)
            for chart_name, chart_path in chart_paths.items():
                if os.path.exists(chart_path):
                    doc.add_picture(chart_path, width=Inches(6))
                    doc.add_paragraph(f"*{chart_name.replace('_', ' ').title()}*")
        
        # Kết luận và khuyến nghị
        doc.add_heading("Kết Luận & Khuyến Nghị", level=1)
        doc.add_paragraph("Dựa trên phân tích dữ liệu, chúng tôi đưa ra các khuyến nghị sau:")
        doc.add_paragraph("• Tập trung vào các sản phẩm bán chạy")
        doc.add_paragraph("• Tối ưu hóa thời gian phục vụ trong giờ cao điểm")
        doc.add_paragraph("• Phát triển chiến lược marketing cho các khung giờ thấp")
        
        # Lưu báo cáo
        today_str = datetime.now().strftime('%d-%m-%Y')
        report_filename = f"Báo cáo ngày {today_str}.docx"

        report_path = os.path.join(REPORTS_DIR, report_filename)
        doc.save(report_path)
        
        print("✅ Tạo báo cáo Word hoàn tất.")
        return {"report_path": report_path}
        
    except Exception as e:
        print(f"❌ Lỗi khi tạo báo cáo: {e}")
        return {"report_path": ""}

def send_email_node(state: AnalysisState) -> Dict[str, Any]:
    """Node để gửi email báo cáo."""
    print("📧 Đang gửi email báo cáo...")
    report_path = state.get("report_path", "")
    analysis = state.get("analysis", "")
    
    try:
        # Cấu hình email (cần thiết lập trong .env)
        smtp_server = os.getenv("SMTP_SERVER", "smtp.gmail.com")
        smtp_port = int(os.getenv("SMTP_PORT", "587"))
        email_user = os.getenv("EMAIL_USER", "")
        email_password = os.getenv("EMAIL_PASSWORD", "")
        recipient_email = os.getenv("RECIPIENT_EMAIL", "")
        
        if not all([email_user, email_password, recipient_email]):
            print("⚠️ Thiếu thông tin email, bỏ qua việc gửi email.")
            return {}
        
        # Tạo email
        msg = MIMEMultipart()
        msg['From'] = email_user
        msg['To'] = recipient_email
        msg['Subject'] = f"Báo Cáo Phân Tích F&B - {datetime.now().strftime('%d/%m/%Y')}"
        
        # Nội dung email
        body = f"""
        Xin chào,
        
        Đính kèm báo cáo phân tích dữ liệu F&B được tạo tự động.
        
        Tóm tắt phân tích:
        {analysis[:500]}...
        
        Trân trọng,
        Hệ thống phân tích tự động
        """
        
        msg.attach(MIMEText(body, 'plain', 'utf-8'))
        
        # Đính kèm file báo cáo
        if report_path and os.path.exists(report_path):
            with open(report_path, "rb") as attachment:
                part = MIMEBase('application', 'octet-stream')
                part.set_payload(attachment.read())
            
            encoders.encode_base64(part)
            part.add_header(
                'Content-Disposition',
                f'attachment; filename= {os.path.basename(report_path)}'
            )
            msg.attach(part)
        
        # Gửi email
        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()
        server.login(email_user, email_password)
        text = msg.as_string()
        server.sendmail(email_user, recipient_email, text)
        server.quit()
        
        print("✅ Gửi email thành công.")
        return {}
        
    except Exception as e:
        print(f"❌ Lỗi khi gửi email: {e}")
        return {}

# --- Xây dựng Graph hoàn chỉnh ---

def build_analysis_graph() -> Any:
    graph = StateGraph(AnalysisState)
    
    # Thêm các node
    graph.add_node("preprocess", preprocess_node)
    graph.add_node("retrieve_knowledge", retrieve_knowledge_node)
    graph.add_node("analyst", llm_analyst_node)
    graph.add_node("critic", llm_critic_node)
    graph.add_node("create_charts", create_charts_node)
    graph.add_node("create_report", create_report_node)
    graph.add_node("send_email", send_email_node)
    
    # Thiết lập quy trình
    graph.set_entry_point("preprocess")
    graph.add_edge("preprocess", "retrieve_knowledge")
    graph.add_edge("retrieve_knowledge", "analyst")
    graph.add_edge("analyst", "critic")
    
    # Thêm cạnh điều kiện
    graph.add_conditional_edges(
        "critic",
        should_continue,
        {
            "reflect": "analyst", # Nếu cần cải thiện, quay lại bước phân tích
            "end": "create_charts"  # Nếu OK, đi tiếp
        }
    )
    
    graph.add_edge("create_charts", "create_report")
    graph.add_edge("create_report", "send_email")
    graph.add_edge("send_email", END)
    
    return graph.compile()

def main(payload: Optional[Dict[str, Any]] = None):
    print("🚀 Bắt đầu quy trình phân tích báo cáo tự động (phiên bản nâng cao)...")
    payload = fetch_data("2025-08-17", "2025-08-22")
    app = build_analysis_graph()
    
    question = "Phân tích dữ liệu kinh doanh và đưa ra các nhận định quan trọng về hiệu suất sản phẩm và xu hướng theo thời gian."
    
    # Chạy quy trình
    if payload is not None:
        app.invoke({"question": question, "payload": payload})
    else:
        app.invoke({"question": question})

    print("\n🎉 Quy trình đã hoàn tất thành công!")

if __name__ == "__main__":
    print("--- CHẠY THỬ NGHIỆM SCRIPT VỚI DỮ LIỆU THẬT ---")

    # ▼▼▼ PASTE YOUR REAL JSON DATA HERE ▼▼▼
    sample_payload_from_n8n = {
      "result": {
        "data": [
          {
            "id": "real_order_001",
            "calcTotalMoney": "150000",
            "createdDateTime": "2025-08-22T14:10:00Z",
            "products": [
              { "productName": "Trà Sữa Trân Châu", "price": "50000", "quantity": "3" }
            ]
          },
          {
            "id": "real_order_002",
            "calcTotalMoney": "45000",
            "createdDateTime": "2025-08-21T18:45:00Z",
            "products": [
              { "productName": "Cà Phê Đen", "price": "20000", "quantity": "1" },
              { "productName": "Bánh Tiramisu", "price": "25000", "quantity": "1" }
            ]
          }
          # ... and so on for all your data
        ]
      }
    }
    # ▲▲▲ PASTE YOUR REAL JSON DATA HERE ▲▲▲

    # The script will run with the data you just pasted
    main(sample_payload_from_n8n)
